import streamlit as st
from io import BytesIO

# Check if python-docx is installed
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    st.error("⚠️ python-docx is not installed. Please run: pip install python-docx")

# Page configuration
st.set_page_config(page_title="MEP SD Narrative Editor", layout="wide", page_icon="📋")

# Initialize session state
if 'form_data' not in st.session_state:
    st.session_state.form_data = {
        'building_name': 'XXXXXXXXXX Building',
        'total_tonnage': 'XXX',
        'space_type_a': 'XXX Space Type A',
        'space_type_a_tonnage': 'XXX',
        'space_type_b': 'XXX Space Type B',
        'space_type_b_tonnage': 'XXX',
        'space_type_c': 'XXX Space Type C',
        'space_type_c_tonnage': 'XXX',
        'hvac_system_type': 'split_dx',
        'cooling_tower_gpm': 'XXXX',
        'chiller_tonnage': 'XXX',
        'vrf_tonnage': 'XXX',
        'outside_air_cfm': 'XXXXX',
        'parking_ventilation': 'mechanical',
        'domestic_water_size': '3"',
        'service_size': '2000',
        'generator_required': True,
        'lightning_protection': False,
        'solar_panels': False,
    }

def create_word_document(data):
    """Generate Word document from form data"""
    if not DOCX_AVAILABLE:
        return None
        
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading(data['building_name'], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_heading('MEP Schematic Design Narrative', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Introduction
    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        'Mechanical, Electrical, Plumbing, & fire protection engineering design approach'
    )
    intro_run.bold = True
    
    doc.add_paragraph(
        f'The following scope outlines the design approach that shall be used for the '
        f'{data["building_name"]} building project for the Mechanical, Electrical, '
        f'Plumbing, & Fire Protection, Systems.'
    )
    
    # MECHANICAL SECTION
    doc.add_heading('MECHANICAL', level=1)
    
    # A. Codes and Standards
    doc.add_heading('A. Codes and Standards:', level=2)
    para = doc.add_paragraph('The building shall be provided with systems in accordance to:')
    codes = [
        '2020 Florida Building Code',
        '2020 Florida Mechanical Code',
        '2020 Florida Energy Code',
        '2020 Florida Plumbing Code',
        '2020 Florida Fuel Gas Code',
        'NFPA 88A, 92, 96, 101'
    ]
    for code in codes:
        doc.add_paragraph(code, style='List Bullet')
    
    # B. Design Conditions
    doc.add_heading('B. Design Conditions:', level=2)
    conditions = [
        'Summer Outside: 92°F DB/ 79°F WB',
        'Winter Outside: 43°F',
        'Summer Inside: 75°F DB/ 55%RH',
        'Winter Inside: 70°F'
    ]
    for condition in conditions:
        doc.add_paragraph(condition, style='List Bullet')
    
    # C. Building Envelope Criteria
    doc.add_heading('C. Building Envelope Criteria:', level=2)
    doc.add_paragraph(
        'Cooling/heating load calculations are currently based on the 2020 Florida Building Code, '
        'Energy Conservation, prescriptive minimums for Climate Zone 2A (Pinellas, Florida):'
    )
    envelope_items = [
        'Climate Zone (Zone 2A)',
        'Built-up Roof: R-25 continuous, U-0.039',
        'Walls (metal frame): R-13.0 + R-6.5ci, U-0.079',
        'Walls (concrete): R-5.7ci, U-0.151',
        'Mass Floors over unconditioned space: R-6.3ci, U-0.107',
        'Fenestration (glazing): Solar Heat Gain Coefficient (SHGC)-0.25, U-0.5'
    ]
    for item in envelope_items:
        doc.add_paragraph(item, style='List Bullet 2')
    
    # G. Building HVAC Systems
    doc.add_heading('G. Building HVAC Systems:', level=2)
    doc.add_paragraph(f'The total building tonnage is estimated at {data["total_tonnage"]} total tons:')
    doc.add_paragraph(f'{data["space_type_a_tonnage"]} Total tons for {data["space_type_a"]}:', style='List Bullet')
    doc.add_paragraph(f'{data["space_type_b_tonnage"]} Total tons for {data["space_type_b"]}:', style='List Bullet')
    doc.add_paragraph(f'{data["space_type_c_tonnage"]} Total tons for {data["space_type_c"]}:', style='List Bullet')
    
    # HVAC System Type
    hvac_systems = {
        'split_dx': 'H. Split Dx Systems:',
        'condenser_water': 'Condenser Water System:',
        'chilled_water': 'Chilled water System with Cooling Towers:',
        'vrf': 'Variable Refrigerant Flow System:'
    }
    
    doc.add_heading(hvac_systems.get(data['hvac_system_type'], 'H. Split Dx Systems:'), level=2)
    
    if data['hvac_system_type'] == 'split_dx':
        split_dx_content = [
            'The HVAC system will be made up of multiple split dx heat pump systems ranging from 1.5 tons to 5 tons.',
            'All condensing units shall be mounted on X" roof curb (or) roof support rails mounted a min. 18" off roof as per the Florida building code chapter 15 table 1510.10.',
            'All condensers shall be coated with a min. 4000 hr salt resistance coating.',
            'All condensers shall be hurricane strapped to structure.',
            'All refrigerant piping located on the exterior of the building will be provided with a UV & corrosion resistant coating.',
            'All refrigerant piping located underground shall be installed in PVC carrier piping.',
        ]
        for content in split_dx_content:
            doc.add_paragraph(content, style='List Bullet')
    elif data['hvac_system_type'] == 'vrf':
        vrf_content = [
            f'The Variable Refrigerant Flow system shall serve XXX Areas with an estimated total tonnage of {data.get("vrf_tonnage", "XXX")} tons.',
            'Variable refrigerant flow systems shall be Heat recovery or Heat Pump.',
            'Branch Controllers shall be provided throughout the system in accordance with the manufacturers recommendations.',
            'All condensers shall be coated with a min. 4000 hr salt resistance coating.',
        ]
        for content in vrf_content:
            doc.add_paragraph(content, style='List Bullet')
    
    # Outside Air System
    doc.add_heading('Outside Air System:', level=2)
    doc.add_paragraph(
        f'The outside air will be provided via XXX roof top units with energy recovery sized at {data["outside_air_cfm"]} cfm each.'
    )
    oa_content = [
        'The outside airs will be packaged DX (or) Chilled water (or) Condenser water.',
        'Provide variable frequency drives for outside air fans. VFDs shall be integral to unit.',
        'Outside air units shall be provided with 8 row coils. Coils shall be selected at no more than 450 FPM.',
    ]
    for content in oa_content:
        doc.add_paragraph(content, style='List Bullet')
    
    # Building Management System
    doc.add_heading('Building Management System:', level=2)
    bms_content = [
        'The building will be provided with a building management system (BMS) that will connect and monitor all the tenant spaces.',
        'The BACnet control system shall be Native level 5 BACnet capable.',
        'The BMS will be located on the ground floor in the engineer\'s office.',
        'Afterhours cooling shall be via a ticketed system that coordinates with the building engineer\'s staff.',
    ]
    for content in bms_content:
        doc.add_paragraph(content, style='List Bullet')
    
    # PLUMBING SECTION
    doc.add_page_break()
    doc.add_heading('PLUMBING', level=1)
    
    doc.add_heading('A. Codes and Standards:', level=2)
    doc.add_paragraph('The building shall be provided with systems in accordance with:')
    plumb_codes = [
        '2017 Florida Building Code',
        '2017 Florida Mechanical Code',
        '2017 Florida Energy Code',
        '2017 Florida Plumbing Code',
        '2017 Florida Fuel Gas Code',
        'Americans with Disabilities Act (ADA)'
    ]
    for code in plumb_codes:
        doc.add_paragraph(code, style='List Bullet')
    
    doc.add_heading('B. Domestic Water:', level=2)
    doc.add_paragraph(
        f'A new domestic water service shall be extended and brought into each building from the existing main. '
        f'The building domestic water shall be {data["domestic_water_size"]} in size and enter the building '
        f'in the 1st level domestic water and fire pump room. The meter and backflow preventer will be '
        f'located outside the building by civil.'
    )
    
    doc.add_heading('C. Plumbing Fixtures:', level=2)
    doc.add_paragraph('Public and General Fixtures:', style='Heading 3')
    fixture_content = [
        'Fixtures shall be ADA compliant where required.',
        'Lavatories shall be equipped with 0.4 gallons per minute maximum sensor activated electronic faucets.',
        'Urinals shall be wall hung, 0.85 gallons per flush (Pint Flush) maximum.',
        'Water closets shall be wall hung, 1.28 maximum gallons per flush.',
        'Floor drains shall be provided in mechanical rooms and in all toilet rooms with more than one water closet.',
    ]
    for content in fixture_content:
        doc.add_paragraph(content, style='List Bullet')
    
    # FIRE PROTECTION SECTION
    doc.add_page_break()
    doc.add_heading('FIRE PROTECTION', level=1)
    
    doc.add_heading('A. Codes and Standards:', level=2)
    doc.add_paragraph(
        'Fire suppressions systems shall be designed and installed to meet the requirements '
        'of all applicable state and local codes as well as required industry standards as follows:'
    )
    fp_codes = [
        '2020 Florida Building Code (FBC)',
        '2020 Florida Fire Prevention Code (FFPC)',
        'NFPA 13 v.2016 -- Standard for the Installation of Sprinkler Systems',
        'NFPA 14 v.2016 -- Standard for the Installation of Standpipe and Hose Systems',
        'NFPA 20 v.2016 -- Standard for the Installation of Stationary Pumps for Fire Protection'
    ]
    for code in fp_codes:
        doc.add_paragraph(code, style='List Bullet')
    
    doc.add_heading('B. Automatic Sprinkler/Standpipe Systems:', level=2)
    sprinkler_content = [
        'Automatic sprinkler systems shall be provided throughout each building as required by FBC 403.3.',
        'A standpipe system shall be provided in each building as required by FBC 403.4.3 and FBC 905.3.',
        'Fire water shall be supplied to each building from the municipal water main.',
        'Each building shall be fully sprinkled throughout with a hydraulically calculated wet pipe automatic sprinkler system.',
    ]
    for content in sprinkler_content:
        doc.add_paragraph(content, style='List Bullet')
    
    # ELECTRICAL SECTION
    doc.add_page_break()
    doc.add_heading('ELECTRICAL', level=1)
    
    doc.add_heading('A. Codes and Standards:', level=2)
    doc.add_paragraph('2017 National Electrical Code NEC (NFPA 70)', style='List Bullet')
    
    doc.add_heading('B. Power Distribution:', level=2)
    doc.add_paragraph(
        f'A new electrical service from the utility Duke Energy will be designed to provide power to the building. '
        f'A Duke owned pad-mount transformer will be installed outside and an underground service lateral will be '
        f'supplied to the Main Electrical Room on the first floor.'
    )
    doc.add_paragraph(
        f'It is anticipated the service lateral will terminate in a {data["service_size"]}A switchboard. '
        f'Power will be distributed from the switchboard to metering equipment and individual tenant panelboards.'
    )
    
    power_content = [
        'The main devices will have ground fault protection.',
        'Power will be distributed at 480 volts for mechanical equipment and 277 volts for lighting.',
        'Power will be stepped down via energy saving, dry type transformers to 208/120 volts.',
        'Power will be distributed at 208V single phase to each residential unit.',
    ]
    for content in power_content:
        doc.add_paragraph(content, style='List Bullet')
    
    if data['generator_required']:
        doc.add_heading('C. Emergency Power Distribution:', level=2)
        gen_content = [
            'An Engine Generator Set for emergency power will be provided for the building.',
            'Emergency power will be distributed via circuit breaker-type distribution panels.',
            'Generator will be as manufactured by Caterpillar, Cummins, MTU.',
            'Automatic transfer switches will be Zenith, Asco, or equal.',
            'All elevators will be connected to the generator.',
        ]
        for content in gen_content:
            doc.add_paragraph(content, style='List Bullet')
    
    doc.add_heading('D. Lighting System:', level=2)
    doc.add_paragraph('Illumination Levels (average maintained):')
    lighting_levels = [
        'Lobbies and Corridors: 20 fc',
        'Restrooms: 25 fc',
        'Storage Rooms: 20 fc',
        'Mechanical/Electrical Rooms: 25 fc',
        'Circulation Stairs: 10 fc',
    ]
    for level in lighting_levels:
        doc.add_paragraph(level, style='List Bullet')
    
    # SOLAR PANELS (if selected)
    if data['solar_panels']:
        doc.add_page_break()
        doc.add_heading('ROOF TOP SOLAR PANEL SYSTEM (ADD ALTERNATE)', level=1)
        solar_content = [
            'The proposed building will provide an option for roof mounted solar.',
            'The roof mounted solar PV system will consist of approximately 230 REC380AA panels that are 380 watts a piece with a 21.7% efficiency.',
            'The total size of the system would be 87 kW.',
            'Each panel will have an Enphase IQ7+ microinverter to convert from DC power to AC 208/3.',
            'The solar system shall tie into the building incoming power from TECO with a net meter installed on site.',
        ]
        for content in solar_content:
            doc.add_paragraph(content, style='List Bullet')
    
    # FIRE ALARM SECTION
    doc.add_page_break()
    doc.add_heading('FIRE ALARM', level=1)
    
    doc.add_heading('A. Codes and Standards:', level=2)
    fa_codes = [
        'Florida Fire Prevention Code NFPA 72',
        'The National Fire Alarm and Signaling Code',
        'NFPA 70 The National Electrical Code'
    ]
    for code in fa_codes:
        doc.add_paragraph(code, style='List Bullet')
    
    doc.add_heading('B. Fire Alarm System:', level=2)
    fa_content = [
        'The building will be equipped throughout with an emergency voice alarm fire alarm system.',
        'The main fire alarm control panel will be located in the fire command center.',
        'Manual pull stations will be installed at all exits.',
        'Smoke detection will be provided in mechanical, electrical, elevator rooms and control rooms.',
        'System smoke detectors functioning as single- and/or multiple-station smoke alarms will be installed in all residential units.',
    ]
    for content in fa_content:
        doc.add_paragraph(content, style='List Bullet')
    
    # Add closing
    doc.add_paragraph()
    doc.add_paragraph()
    closing = doc.add_paragraph('End of Narrative')
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing.runs[0].bold = True
    
    return doc

def get_download_link(doc):
    """Generate download link for Word document"""
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

# Main App Layout
st.title("📋 MEP Schematic Design Narrative Editor")
st.markdown("---")

# Show installation instructions if library not available
if not DOCX_AVAILABLE:
    st.warning("""
    ### 📦 Installation Required
    
    To use the Word export feature, please install the required library:
    
    ```bash
    pip install python-docx
    ```
    
    Then restart the Streamlit app.
    """)

# Sidebar for key inputs
with st.sidebar:
    st.header("🏢 Project Information")
    st.session_state.form_data['building_name'] = st.text_input(
        "Building Name",
        value=st.session_state.form_data['building_name']
    )
    
    st.markdown("### System Selections")
    st.session_state.form_data['hvac_system_type'] = st.selectbox(
        "HVAC System Type",
        options=['split_dx', 'condenser_water', 'chilled_water', 'vrf'],
        format_func=lambda x: {
            'split_dx': 'Split DX Systems',
            'condenser_water': 'Condenser Water System',
            'chilled_water': 'Chilled Water System',
            'vrf': 'Variable Refrigerant Flow'
        }[x],
        index=['split_dx', 'condenser_water', 'chilled_water', 'vrf'].index(
            st.session_state.form_data['hvac_system_type']
        )
    )
    
    st.session_state.form_data['generator_required'] = st.checkbox(
        "Emergency Generator Required",
        value=st.session_state.form_data['generator_required']
    )
    
    st.session_state.form_data['solar_panels'] = st.checkbox(
        "Include Solar Panels (Add Alternate)",
        value=st.session_state.form_data['solar_panels']
    )
    
    st.session_state.form_data['lightning_protection'] = st.checkbox(
        "Lightning Protection",
        value=st.session_state.form_data['lightning_protection']
    )
    
    st.markdown("---")
    
    # Export button
    if DOCX_AVAILABLE:
        if st.button("📥 Generate & Download Word Document", type="primary", use_container_width=True):
            with st.spinner("Generating document..."):
                doc = create_word_document(st.session_state.form_data)
                doc_bytes = get_download_link(doc)
                
                st.download_button(
                    label="⬇️ Download Word Document",
                    data=doc_bytes,
                    file_name=f"{st.session_state.form_data['building_name'].replace(' ', '_')}_MEP_Narrative.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                st.success("✅ Document generated successfully!")
    else:
        st.button("📥 Generate Document (Disabled)", disabled=True, use_container_width=True)
        st.caption("Install python-docx to enable")

# Main content area with tabs
tab1, tab2, tab3, tab4 = st.tabs(["🔧 Mechanical", "💧 Plumbing", "⚡ Electrical", "🔥 Fire Protection"])

with tab1:
    st.header("Mechanical Systems")
    
    with st.expander("📊 Building HVAC Systems", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            st.session_state.form_data['total_tonnage'] = st.text_input(
                "Total Building Tonnage",
                value=st.session_state.form_data['total_tonnage']
            )
            st.session_state.form_data['space_type_a'] = st.text_input(
                "Space Type A Description",
                value=st.session_state.form_data['space_type_a']
            )
            st.session_state.form_data['space_type_a_tonnage'] = st.text_input(
                "Space Type A Tonnage",
                value=st.session_state.form_data['space_type_a_tonnage']
            )
        with col2:
            st.session_state.form_data['space_type_b'] = st.text_input(
                "Space Type B Description",
                value=st.session_state.form_data['space_type_b']
            )
            st.session_state.form_data['space_type_b_tonnage'] = st.text_input(
                "Space Type B Tonnage",
                value=st.session_state.form_data['space_type_b_tonnage']
            )
            st.session_state.form_data['space_type_c'] = st.text_input(
                "Space Type C Description",
                value=st.session_state.form_data['space_type_c']
            )
            st.session_state.form_data['space_type_c_tonnage'] = st.text_input(
                "Space Type C Tonnage",
                value=st.session_state.form_data['space_type_c_tonnage']
            )
    
    with st.expander("🌡️ Design Conditions"):
        st.info("Summer Outside: 92°F DB/ 79°F WB | Winter Outside: 43°F | Summer Inside: 75°F DB/ 55%RH | Winter Inside: 70°F")
    
    with st.expander("🏗️ Building Envelope Criteria"):
        st.info("Based on 2020 Florida Building Code, Energy Conservation, prescriptive minimums for Climate Zone 2A")
        st.write("**Key Parameters:**")
        st.write("- Built-up Roof: R-25 continuous, U-0.039")
        st.write("- Walls (metal frame): R-13.0 + R-6.5ci, U-0.079")
        st.write("- Fenestration (glazing): SHGC-0.25, U-0.5")
    
    with st.expander("💨 Outside Air System"):
        st.session_state.form_data['outside_air_cfm'] = st.text_input(
            "Outside Air CFM",
            value=st.session_state.form_data['outside_air_cfm']
        )
        st.write("- Outside air provided via roof top units with energy recovery")
        st.write("- Variable frequency drives for outside air fans")
        st.write("- Sectional enthalpy wheel for energy recovery")
    
    with st.expander("🅿️ Parking Ventilation"):
        st.session_state.form_data['parking_ventilation'] = st.radio(
            "Parking Ventilation Type",
            options=['mechanical', 'open_air'],
            format_func=lambda x: 'Mechanical Ventilation (NFPA 88A)' if x == 'mechanical' else 'Open Air Garage',
            index=0 if st.session_state.form_data['parking_ventilation'] == 'mechanical' else 1
        )
        if st.session_state.form_data['parking_ventilation'] == 'mechanical':
            st.write("- Ventilation at 1 cfm per square foot")
            st.write("- Activated by CO/NO2 sensors")

with tab2:
    st.header("Plumbing Systems")
    
    with st.expander("💧 Domestic Water", expanded=True):
        st.session_state.form_data['domestic_water_size'] = st.text_input(
            "Domestic Water Service Size",
            value=st.session_state.form_data['domestic_water_size']
        )
        st.write("- New service from existing main")
        st.write("- Enters building in 1st level domestic water and fire pump room")
        st.write("- Meter and backflow preventer located outside by civil")
        st.write("- Packaged duplex domestic water booster pump required")
    
    with st.expander("🚽 Plumbing Fixtures"):
        st.write("**Public and General Fixtures**")
        st.write("- Fixtures shall be ADA compliant where required")
        st.write("- Lavatories: 0.4 gallons per minute maximum sensor activated")
        st.write("- Urinals: 0.85 gallons per flush (Pint Flush) maximum")
        st.write("- Water closets: 1.28 maximum gallons per flush")
        st.write("- Floor drains in mechanical rooms and toilet rooms")
    
    with st.expander("⚙️ Plumbing Equipment"):
        st.write("- Packaged duplex domestic water booster pump (25 psi minimum at 4th floor)")
        st.write("- Water sub-meters for each retail tenant unit and office levels")
        st.write("- Simplex submersible sump pump for each elevator pit")
        st.write("- Traffic rated grease interceptor (1,500 gallons)")
        st.write("- Reduced pressure backflow preventer for public water supply protection")
    
    with st.expander("🌧️ Storm Drainage"):
        st.write("- Roof drains with overflow drains for secondary drainage")
        st.write("- Underground piping: PVC DWV pipe and fittings")
        st.write("- Rainwater collection directed to storm water retention vault")
        st.write("- Duplex skid with grey water filter system for irrigation reuse")

with tab3:
    st.header("Electrical Systems")
    
    with st.expander("⚡ Power Distribution", expanded=True):
        st.session_state.form_data['service_size'] = st.text_input(
            "Service Size (Amps)",
            value=st.session_state.form_data['service_size']
        )
        st.write(f"**Main Service:** {st.session_state.form_data['service_size']}A @ 480V/277VAC 3ɸ")
        st.write("- Duke Energy pad-mount transformer")
        st.write("- 480V for mechanical equipment")
        st.write("- 277V for lighting")
        st.write("- 208/120V for receptacles and residential units")
        st.write("- Each condo: dedicated 208/120V single phase panel (minimum 125A)")
    
    with st.expander("💡 Lighting System"):
        st.write("**Illumination Levels (average maintained):**")
        st.write("- Lobbies and Corridors: 20 fc")
        st.write("- Restrooms: 25 fc")
        st.write("- Storage Rooms: 20 fc")
        st.write("- Mechanical/Electrical Rooms: 25 fc")
        st.write("- Circulation Stairs: 10 fc")
        st.write("- Shell spaces: 1 fc minimum along egress path")
        st.write("\n**Control:**")
        st.write("- Dual technology occupancy sensors in storage and common areas")
        st.write("- Daylight harvesting per ASHRAE 90.1")
        st.write("- LED exit signs")
    
    with st.expander("🔌 Wiring Devices"):
        st.write("**Specifications:**")
        st.write("- Specification grade: Hubbell, Pass & Seymour, Cooper")
        st.write("- GFCI receptacles per NEC requirements")
        st.write("- Break resistant 'Noryl' cover plates")
        st.write("- 20A, 120V branch circuits: 1500W maximum")
        st.write("- 20A, 277V branch circuits: 4200W maximum")
    
    if st.session_state.form_data['generator_required']:
        with st.expander("🔋 Emergency Power Distribution"):
            st.write("**Generator System:**")
            st.write("- Manufacturer: Caterpillar, Cummins, MTU")
            st.write("- Automatic transfer switches: Zenith, Asco")
            st.write("\n**Loads Served:**")
            st.write("- Life safety loads")
            st.write("- Legally required standby loads")
            st.write("- Optional standby systems")
            st.write("- All elevators (1 at a time operation)")
            st.write("- Fire pump (dedicated ATS by manufacturer)")
            st.write("- Stair and hoist way pressurization")

with tab4:
    st.header("Fire Protection & Fire Alarm Systems")
    
    with st.expander("🚿 Automatic Sprinkler/Standpipe Systems", expanded=True):
        st.write("**Sprinkler System:**")
        st.write("- Automatic sprinkler systems throughout per FBC 403.3")
        st.write("- Hydraulically calculated wet pipe system per NFPA 13")
        st.write("- Zoned per floor minimum (FBC 903.4.3)")
        st.write("\n**Standpipe System:**")
        st.write("- Class I standpipes in each egress stairwell")
        st.write("- 2-1/2\" fire hose connections per FBC 905.4")
        st.write("- Automatic wet standpipe per NFPA 14")
        st.write("\n**Fire Pump:**")
        st.write("- Electric fire pump per NFPA 20")
        st.write("- 100 psi residual pressure to most remote hose connection")
